import os
import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import logging

logger = logging.getLogger(__name__)

try:
    from core.ats_optimizer import (
        extract_ats_keywords,
        optimize_for_ats,
        calculate_ats_score,
        clean_for_ats,
        generate_ats_tips,
    )

    ATS_AVAILABLE = True
except ImportError:
    ATS_AVAILABLE = False
    logger.warning("ATS optimizer not available")

CV_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "CV-trabajp")


TECH_KEYWORDS = {
    "python": [
        "Python",
        "Django",
        "Flask",
        "FastAPI",
        "Pandas",
        "NumPy",
        "API",
        "Automation",
        "Scripting",
    ],
    "javascript": [
        "JavaScript",
        "React",
        "Vue",
        "Node.js",
        "Express",
        "TypeScript",
        "Frontend",
        "jQuery",
    ],
    "java": ["Java", "Spring", "Hibernate", "JVM", "Microservices", "REST API"],
    "sql": ["SQL", "PostgreSQL", "MySQL", "MongoDB", "Database", "Queries", "NoSQL"],
    "aws": [
        "AWS",
        "Amazon Web Services",
        "EC2",
        "S3",
        "Lambda",
        "Cloud",
        "Azure",
        "GCP",
    ],
    "docker": [
        "Docker",
        "Kubernetes",
        "Container",
        "DevOps",
        "CI/CD",
        "Jenkins",
        "GitHub Actions",
    ],
    "ai": [
        "AI",
        "Artificial Intelligence",
        "Machine Learning",
        "ChatGPT",
        "GPT",
        "LLM",
        "NLP",
        "Automation",
    ],
    "data": ["Data Analysis", "Excel", "Tableau", "Power BI", "Analytics", "Reporting"],
    "remote": ["Remote", "Virtual", "Home Office", "Telecommuting", "Async"],
    "soft_skills": [
        "Communication",
        "Teamwork",
        "Problem Solving",
        "Leadership",
        "Time Management",
    ],
}


def extract_job_keywords(job_description):
    if not job_description:
        return []

    desc_lower = job_description.lower()
    found_keywords = []

    for category, keywords in TECH_KEYWORDS.items():
        for kw in keywords:
            if kw.lower() in desc_lower:
                found_keywords.append(kw)

    tech_patterns = [
        r"\b[A-Z][a-z]+\s*\d+\.\d+\b",
        r"\b\w+\+\b",
        r"\b\w+#\d+\b",
        r"\b[A-Z]{2,}\b",
    ]

    for pattern in tech_patterns:
        matches = re.findall(pattern, job_description)
        found_keywords.extend(matches)

    return list(set(found_keywords))[:20]


def extract_missing_keywords(job_description, cv_content):
    if not cv_content or not job_description:
        return []

    job_kw = set(extract_job_keywords(job_description))
    cv_kw = set(kw.lower() for kw in extract_job_keywords(cv_content))

    missing = job_kw - cv_kw
    return list(missing)


def adapt_cv_for_job(cv_path, job_data, output_folder="data/cvs"):
    job_title = job_data.get("title", "")
    company = job_data.get("company", "Unknown")
    job_description = job_data.get("description", "")

    os.makedirs(output_folder, exist_ok=True)

    try:
        doc = Document(cv_path)
        adapted_doc = Document()

        for para in doc.paragraphs:
            new_text = para.text

            if job_description:
                keywords = extract_job_keywords(job_description)
                for kw in keywords:
                    if kw.lower() not in new_text.lower():
                        if (
                            "skill" in para.text.lower()
                            or "ability" in para.text.lower()
                        ):
                            new_text = new_text + f", {kw}"

            if new_text:
                adapted_para = adapted_doc.add_paragraph(new_text)
                try:
                    adapted_para.runs[0].bold = (
                        para.runs[0].bold if para.runs else False
                    )
                    adapted_para.runs[0].font.size = (
                        para.runs[0].font.size if para.runs else None
                    )
                except:
                    pass

        filename = (
            f"CV_{company.replace(' ', '_')}_{job_title.replace(' ', '_')[:20]}.docx"
        )
        filepath = os.path.join(output_folder, filename)
        adapted_doc.save(filepath)

        logger.info(f"Adapted CV saved to {filepath}")
        return filepath

    except Exception as e:
        logger.error(f"Error adapting CV: {e}")
        return None


def create_adapted_cv(cv_manager, job, output_folder="data/cvs"):
    from core.cv_manager import CVManager

    cv_path = cv_manager.get_recommended_cv(
        job.title, remote=(job.remote_type == "remote")
    )

    if not cv_path:
        return None

    job_data = {
        "title": job.title,
        "company": job.company,
        "description": job.description,
    }

    adapted_path = adapt_cv_for_job(cv_path["path"], job_data, output_folder)

    return adapted_path


def generate_adapted_cover_letter(cv_content, job_data):
    from core.cover_letter import generate_cover_letter

    from datetime import datetime

    personal_info = extract_personal_info_from_cv(cv_content)

    cover_letter = generate_cover_letter(job_data, personal_info)

    return cover_letter


def extract_personal_info_from_cv(cv_content):
    info = {
        "name": "",
        "email": "",
        "phone": "",
        "location": "",
        "experience_summary": "",
    }

    lines = cv_content.split("\n")[:30]

    for i, line in enumerate(lines):
        line_lower = line.lower()

        if "@" in line and "email" not in line_lower:
            info["email"] = line.strip()
        elif any(x in line_lower for x in ["phone", "tel", "+"]) and not info["phone"]:
            info["phone"] = line.strip()
        elif (
            any(x in line_lower for x in ["location", "bolivia", "peru", "argentina"])
            and not info["location"]
        ):
            info["location"] = line.strip()
        elif i < 5 and not info["name"]:
            if len(line.strip()) < 50 and line.strip():
                info["name"] = line.strip()

    if not info["name"]:
        info["name"] = "Candidate"

    return info


def adapt_and_queue_application(job_id):
    from models.job import Job
    from models.database import db
    from models.application import Application

    job = Job.query.get(job_id)
    if not job:
        return None, "Job not found"

    from core.cv_manager import CVManager

    cv_manager = CVManager()

    recommended_cv = cv_manager.get_recommended_cv(
        job.title, remote=(job.remote_type == "remote")
    )
    if not recommended_cv:
        return None, "No suitable CV found"

    cv_content = cv_manager.extract_content_from_docx(recommended_cv["path"])

    job_data = {
        "title": job.title,
        "company": job.company,
        "description": job.description,
    }

    adapted_cv_path = adapt_cv_for_job(recommended_cv["path"], job_data)

    cover_letter = generate_adapted_cover_letter(cv_content, job_data)

    from core.cover_letter import save_cover_letter

    cover_path = save_cover_letter(cover_letter, job.company, job.title)

    application = Application(
        job_id=job_id,
        status="pending_review",
        needs_review=True,
        cv_path=adapted_cv_path,
        cover_letter_path=cover_path,
    )

    db.session.add(application)
    db.session.commit()

    return application, "Application created with adapted CV"


def batch_adapt_cvs(job_ids):
    results = []
    for job_id in job_ids:
        app, msg = adapt_and_queue_application(job_id)
        results.append(
            {"job_id": job_id, "application": app.id if app else None, "message": msg}
        )

    return results
