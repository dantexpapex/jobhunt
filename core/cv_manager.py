import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import Config
import logging

logger = logging.getLogger(__name__)

CV_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "CV-trabajp")


class CVManager:
    def __init__(self, cv_dir=CV_DIR):
        self.cv_dir = cv_dir
        self.available_cvs = self.list_cvs()

    def list_cvs(self):
        cvs = []
        for file in os.listdir(self.cv_dir):
            if file.endswith(".docx"):
                cvs.append(
                    {
                        "filename": file,
                        "path": os.path.join(self.cv_dir, file),
                        "name": file.replace(".docx", "").replace("_", " ").title(),
                    }
                )
        return cvs

    def get_cv_by_keyword(self, keyword):
        keyword = keyword.lower()
        for cv in self.available_cvs:
            if keyword in cv["filename"].lower():
                return cv
        return self.available_cvs[0] if self.available_cvs else None

    def get_recommended_cv(self, job_title, remote=True):
        title = job_title.lower()

        if "automation" in title or "auto" in title:
            return self.get_cv_by_keyword("automation")
        elif "web developer" in title or "frontend" in title or "front-end" in title:
            return self.get_cv_by_keyword("web developer")
        elif "junior" in title or "entry" in title:
            return self.get_cv_by_keyword("junior")
        elif "assistant" in title or "virtual" in title:
            return self.get_cv_by_keyword("asistente")
        elif "portug" in title:
            return self.get_cv_by_keyword("portug")
        elif remote and "remote" in title:
            return self.get_cv_by_keyword("remoto")

        return self.get_cv_by_keyword("profesional")

    def extract_content_from_docx(self, docx_path):
        try:
            doc = Document(docx_path)
            content = []

            for para in doc.paragraphs:
                if para.text.strip():
                    content.append(para.text)

            return "\n".join(content)
        except Exception as e:
            logger.error(f"Error reading CV: {e}")
            return ""

    def extract_sections_from_docx(self, docx_path):
        try:
            doc = Document(docx_path)
            sections = {
                "summary": "",
                "experience": [],
                "education": [],
                "skills": [],
            }

            current_section = None
            for para in doc.paragraphs:
                text = para.text.strip().lower()

                if "experience" in text or "experiencia" in text:
                    current_section = "experience"
                elif "education" in text or "educación" in text or "educacao" in text:
                    current_section = "education"
                elif "skill" in text or "habilidad" in text or "competencia" in text:
                    current_section = "skills"
                elif "summary" in text or "perfil" in text or "objetivo" in text:
                    current_section = "summary"

                if current_section and para.text.strip():
                    if current_section == "experience":
                        sections["experience"].append(para.text)
                    elif current_section == "education":
                        sections["education"].append(para.text)
                    elif current_section == "skills":
                        sections["skills"].append(para.text)
                    elif current_section == "summary":
                        sections["summary"] += para.text + " "

            return sections
        except Exception as e:
            logger.error(f"Error extracting sections: {e}")
            return sections

    def copy_and_adapt_cv(self, source_cv, target_path, job_data):
        try:
            import shutil

            shutil.copy2(source_cv["path"], target_path)
            logger.info(f"Copied CV to {target_path}")
            return target_path
        except Exception as e:
            logger.error(f"Error copying CV: {e}")
            return None


def get_available_cvs():
    manager = CVManager()
    return manager.available_cvs


def get_cv_for_job(job_title, remote=True):
    manager = CVManager()
    return manager.get_recommended_cv(job_title, remote)


def use_cv(cv_filename, job_id=None):
    manager = CVManager()
    cv = manager.get_cv_by_keyword(cv_filename)

    if not cv:
        return None

    from models.database import db
    from models.cv import CV

    existing = CV.query.filter_by(name=cv["name"]).first()
    if existing:
        return existing

    content = manager.extract_content_from_docx(cv["path"])
    sections = manager.extract_sections_from_docx(cv["path"])

    new_cv = CV(
        name=cv["name"], file_path=cv["path"], base_cv=True, content=str(sections)
    )

    db.session.add(new_cv)
    db.session.commit()

    return new_cv


def auto_select_cv(job_id):
    from models.job import Job

    job = Job.query.get(job_id)
    if not job:
        return None

    manager = CVManager()
    cv = manager.get_recommended_cv(job.title, remote=(job.remote_type == "remote"))

    if cv:
        return use_cv(cv["filename"], job_id)

    return None
